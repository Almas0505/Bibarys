#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Скрипт конвертации Markdown в Word документ по ГОСТ
Использует python-docx для создания DOCX файла
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def convert_markdown_to_word(md_file, output_file):
    """Конвертация Markdown файла в Word документ"""
    
    # Создаем новый документ
    doc = Document()
    
    # НАСТРОЙКА ПОЛЕЙ СТРАНИЦЫ (ГОСТ)
    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(3.0)    # Левое: 30 мм
        section.right_margin = Cm(1.0)   # Правое: 10 мм
        section.top_margin = Cm(2.0)     # Верхнее: 20 мм
        section.bottom_margin = Cm(2.5)  # Нижнее: 25 мм
    
    # НАСТРОЙКА ОСНОВНОГО СТИЛЯ (ГОСТ)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.color.rgb = RGBColor(0, 0, 0)
    
    # Межстрочный интервал: одинарный
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.first_line_indent = Cm(1.25)  # Красная строка 1.25 см
    
    # НАСТРОЙКА ЗАГОЛОВКОВ (ГОСТ)
    # Заголовок 1 уровня: ПРОПИСНЫЕ, 14pt, жирный, с абзацного отступа
    heading1_style = doc.styles['Heading 1']
    heading1_style.font.name = 'Times New Roman'
    heading1_style.font.size = Pt(14)
    heading1_style.font.bold = True
    heading1_style.font.color.rgb = RGBColor(0, 0, 0)
    heading1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading1_style.paragraph_format.first_line_indent = Cm(1.25)
    heading1_style.paragraph_format.space_before = Pt(0)
    heading1_style.paragraph_format.space_after = Pt(0)
    
    # Заголовки 2-3 уровня: 14pt, жирный, с абзацного отступа
    for i in range(2, 5):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(0, 0, 0)
        heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading_style.paragraph_format.first_line_indent = Cm(1.25)
        heading_style.paragraph_format.space_before = Pt(0)
        heading_style.paragraph_format.space_after = Pt(0)
    
    # Читаем Markdown файл
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Предобработка: удаляем ссылки, оставляя только текст
    content = remove_markdown_links(content)
    
    # Разбиваем на строки
    lines = content.split('\n')
    
    i = 0
    prev_was_empty = False
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Пропускаем множественные пустые строки
        if not line:
            if not prev_was_empty:
                doc.add_paragraph()
                prev_was_empty = True
            i += 1
            continue
        
        prev_was_empty = False
        
        # Заголовок 1 уровня (# ) - ПРОПИСНЫЕ БУКВЫ
        if line.startswith('# ') and not line.startswith('## '):
            text = clean_text(line[2:].strip())
            text = text.upper()  # ДЕЛАЕМ ВСЕ БУКВЫ ПРОПИСНЫМИ
            p = doc.add_heading(level=1)
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            i += 1
            continue
        
        # Заголовок 2 уровня (## )
        if line.startswith('## ') and not line.startswith('### '):
            text = clean_text(line[3:].strip())
            p = doc.add_heading(level=2)
            add_formatted_run(p, text, is_heading=True)
            i += 1
            continue
        
        # Заголовок 3 уровня (### )
        if line.startswith('### ') and not line.startswith('#### '):
            text = clean_text(line[4:].strip())
            p = doc.add_heading(level=3)
            add_formatted_run(p, text, is_heading=True)
            i += 1
            continue
        
        # Заголовок 4 уровня (#### )
        if line.startswith('#### '):
            text = clean_text(line[5:].strip())
            p = doc.add_heading(level=4)
            add_formatted_run(p, text, is_heading=True)
            i += 1
            continue
        
        # Списки (- или *)
        if line.startswith('- ') or line.startswith('* '):
            text = clean_text(line[2:].strip())
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_run(p, text)
            # Убираем красную строку для списков
            p.paragraph_format.first_line_indent = Cm(0)
            i += 1
            continue
        
        # Нумерованные списки - добавляем номер как часть текста
        if re.match(r'^\d+\.\s', line):
            # Извлекаем номер и текст
            match = re.match(r'^(\d+)\.\s(.+)', line)
            if match:
                number = match.group(1)
                text = clean_text(match.group(2).strip())
                
                # Создаем обычный параграф (НЕ List Number!)
                p = doc.add_paragraph()
                
                # Добавляем номер как часть текста
                num_run = p.add_run(f'{number}. ')
                num_run.font.name = 'Times New Roman'
                num_run.font.size = Pt(14)
                num_run.font.color.rgb = RGBColor(0, 0, 0)
                
                # Добавляем остальной текст с форматированием
                add_formatted_run(p, text)
                
                # Убираем красную строку для списков
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.left_indent = Cm(0)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            i += 1
            continue
        
        # Код блоки (```) - каждая строка = отдельный параграф (работает в Word Online)
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            # Добавляем каждую строку кода как отдельный параграф
            for code_line in code_lines:
                p = doc.add_paragraph()
                # Если строка пустая, добавляем пробел чтобы сохранить пустую строку
                run = p.add_run(code_line if code_line.strip() else ' ')
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
                
                # Убираем красную строку для кода
                p.paragraph_format.first_line_indent = Cm(0)
                # Убираем отступы между строками кода
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                
                # Светло-серая заливка для выделения кода
                set_paragraph_shading(p, 'F5F5F5')
            
            # Добавляем пустую строку после блока кода
            doc.add_paragraph()
            
            i += 1
            continue
        
        # Таблицы (начинаются с |)
        if line.strip().startswith('|') and '|' in line:
            # Собираем все строки таблицы
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            if len(table_lines) >= 2:
                # Удаляем строку-разделитель (вторая строка с ----)
                header_line = table_lines[0]
                data_lines = [l for l in table_lines[2:] if not all(c in '|-: ' for c in l)]
                
                # Парсим заголовки
                headers = [cell.strip() for cell in header_line.split('|') if cell.strip()]
                
                # Создаем таблицу
                if data_lines:
                    num_cols = len(headers)
                    num_rows = len(data_lines) + 1  # +1 для заголовка
                    
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Table Grid'
                    
                    # Заполняем заголовки
                    for col_idx, header_text in enumerate(headers):
                        cell = table.rows[0].cells[col_idx]
                        cell.text = clean_text(header_text)
                        # Жирный шрифт для заголовков
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(14)
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    # Заполняем данные
                    for row_idx, data_line in enumerate(data_lines, start=1):
                        cells_data = [cell.strip() for cell in data_line.split('|') if cell.strip()]
                        for col_idx, cell_text in enumerate(cells_data[:num_cols]):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = clean_text(cell_text)
                            # Форматируем текст ячейки
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(14)
                                    run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    # Добавляем пустую строку после таблицы
                    doc.add_paragraph()
            
            continue
        
        # Горизонтальная линия (---) - пропускаем
        if line.strip() == '---':
            i += 1
            continue
        
        # Обычный параграф
        text = clean_text(line)
        p = doc.add_paragraph()
        add_formatted_run(p, text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        i += 1
    
    # Сохраняем документ
    doc.save(output_file)
    print(f'✅ Документ успешно создан: {output_file}')


def clean_text(text):
    """Очистка текста от markdown разметки"""
    # Удаляем inline код
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text


def remove_markdown_links(text):
    """Удаление markdown ссылок, оставляя только текст"""
    # [текст](url) -> текст
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    return text


def add_formatted_run(paragraph, text, is_heading=False):
    """Добавление текста с форматированием (жирный текст)"""
    
    # Разбиваем текст на части с жирным форматированием
    # Ищем **текст** или __текст__
    pattern = r'(\*\*|__)(.*?)\1'
    
    last_pos = 0
    has_formatting = False
    
    for match in re.finditer(pattern, text):
        has_formatting = True
        # Добавляем текст до жирного
        if match.start() > last_pos:
            normal_text = text[last_pos:match.start()]
            run = paragraph.add_run(normal_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Добавляем жирный текст
        bold_text = match.group(2)
        run = paragraph.add_run(bold_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        last_pos = match.end()
    
    # Если найдено форматирование, добавляем оставшийся текст
    if has_formatting:
        if last_pos < len(text):
            remaining_text = text[last_pos:]
            run = paragraph.add_run(remaining_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)
    else:
        # Если нет форматирования, добавляем весь текст как есть
        run = paragraph.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)


def set_paragraph_shading(paragraph, color):
    """Установка фонового цвета параграфа"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    paragraph._element.get_or_add_pPr().append(shading_elm)


if __name__ == '__main__':
    import sys
    
    # Параметры по умолчанию или из командной строки
    if len(sys.argv) >= 3:
        input_file = sys.argv[1]
        output_file = sys.argv[2]
    else:
        input_file = 'dd.md'
        output_file = 'ДИПЛОМ_ПОЛНЫЙ.docx'
    
    print('=' * 70)
    print('КОНВЕРТАЦИЯ MARKDOWN → WORD (ПО ГОСТ 7.32-2017)')
    print('=' * 70)
    print(f'\n📄 Входной файл:  {input_file}')
    print(f'📝 Выходной файл: {output_file}\n')
    
    print('⚙️  Применяемые настройки:')
    print('   • Поля: Левое 30мм, Правое 10мм, Верх 20мм, Низ 25мм')
    print('   • Основной текст: Times New Roman 14pt, по ширине')
    print('   • Красная строка: 1.25 см')
    print('   • Межстрочный интервал: одинарный (1.0)')
    print('   • Заголовок 1: ПРОПИСНЫЕ БУКВЫ, 14pt, жирный')
    print('   • Заголовки 2-3: 14pt, жирный, с абзацного отступа')
    print('   • Код: Courier New 10pt, серая заливка\n')
    
    convert_markdown_to_word(input_file, output_file)
    
    print('\n✅ Конвертация завершена успешно!')
    print('\n📋 Следующие шаги:')
    print('   1. Откройте документ в Microsoft Word')
    print('   2. Добавьте титульный лист')
    print('   3. Вставьте оглавление (Ссылки → Оглавление)')
    print('   4. Добавьте изображения (СУРЕТ 1-6)')
    print('   5. Добавьте нумерацию страниц (Вставка → Номер страницы)')
    print('   6. Экспортируйте в PDF при необходимости')
    print('=' * 70)
