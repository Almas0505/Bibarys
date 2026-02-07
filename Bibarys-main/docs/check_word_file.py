#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Скрипт для проверки Word документа на ошибки
"""

import re
from docx import Document
from collections import Counter

def check_word_document(filename):
    """Проверка Word документа на различные ошибки"""
    
    print("=" * 80)
    print("АНАЛИЗ WORD ДОКУМЕНТА")
    print("=" * 80)
    print(f"\n📄 Файл: {filename}\n")
    
    try:
        doc = Document(filename)
    except Exception as e:
        print(f"❌ ОШИБКА: Не удалось открыть файл: {e}")
        return
    
    errors = []
    warnings = []
    
    # Статистика
    total_paragraphs = len(doc.paragraphs)
    headings = []
    numbered_lists = []
    code_blocks = []
    
    print("🔍 Анализирую структуру документа...\n")
    
    # Анализ параграфов
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if not text:
            continue
        
        # Проверка заголовков
        if para.style.name.startswith('Heading'):
            headings.append({
                'level': para.style.name,
                'text': text[:60],
                'line': i + 1
            })
            
            # Проверка заголовка 1 уровня на ПРОПИСНЫЕ буквы
            if para.style.name == 'Heading 1':
                if text != text.upper() and not any(char.isdigit() for char in text):
                    errors.append(f"❌ Строка {i+1}: Заголовок 1 уровня не в ПРОПИСНЫХ буквах: '{text[:40]}...'")
        
        # Поиск нумерованных списков
        if re.match(r'^\d+\.\s', text):
            match = re.match(r'^(\d+)\.\s(.+)', text)
            if match:
                num = int(match.group(1))
                numbered_lists.append({
                    'number': num,
                    'text': match.group(2)[:50],
                    'line': i + 1
                })
        
        # Поиск кода (по шрифту Courier New)
        for run in para.runs:
            if run.font.name == 'Courier New':
                if len(run.text.strip()) > 10:
                    code_blocks.append({
                        'text': run.text[:50],
                        'line': i + 1
                    })
                    break
        
        # Проверка на остатки markdown разметки
        if '**' in text or '__' in text:
            errors.append(f"❌ Строка {i+1}: Найдены markdown маркеры жирного текста: '{text[:60]}...'")
        
        if '```' in text:
            errors.append(f"❌ Строка {i+1}: Найдены markdown маркеры кода: '{text[:60]}...'")
        
        if re.search(r'\[.+\]\(.+\)', text):
            errors.append(f"❌ Строка {i+1}: Найдена markdown ссылка: '{text[:60]}...'")
        
        # Проверка на видимые HTML теги
        if re.search(r'<[a-z]+>', text, re.IGNORECASE):
            warnings.append(f"⚠️  Строка {i+1}: Возможно присутствуют HTML теги: '{text[:60]}...'")
    
    # Анализ нумерованных списков
    print("📊 СТАТИСТИКА ДОКУМЕНТА:")
    print(f"   • Всего параграфов: {total_paragraphs}")
    print(f"   • Таблиц: {len(doc.tables)}")
    print(f"   • Заголовков: {len(headings)}")
    print(f"   • Нумерованных элементов: {len(numbered_lists)}")
    print(f"   • Блоков кода: {len(code_blocks)}")
    
    # Проверка структуры заголовков
    print(f"\n📑 СТРУКТУРА ЗАГОЛОВКОВ:")
    heading_levels = Counter([h['level'] for h in headings])
    for level, count in sorted(heading_levels.items()):
        print(f"   • {level}: {count} шт.")
    
    # Показываем первые 5 заголовков
    print(f"\n📋 ПЕРВЫЕ ЗАГОЛОВКИ:")
    for h in headings[:5]:
        print(f"   {h['level']} (строка {h['line']}): {h['text']}")
    
    # Анализ нумерации списка литературы
    if numbered_lists:
        print(f"\n📚 ПРОВЕРКА НУМЕРАЦИИ СПИСКОВ:")
        
        # Ищем список литературы (последний большой нумерованный список)
        # Находим где начинается новая нумерация с 1
        restarts = []
        for i, item in enumerate(numbered_lists):
            if item['number'] == 1:
                restarts.append(i)
        
        if len(restarts) > 0:
            print(f"   • Найдено {len(restarts)} списков, начинающихся с '1.'")
            
            # Проверяем последний список (должен быть список литературы)
            if len(restarts) > 0:
                last_restart = restarts[-1]
                literature_items = numbered_lists[last_restart:]
                
                print(f"   • Последний список (ЛИТЕРАТУРА): {len(literature_items)} элементов")
                print(f"   • Начинается со строки {literature_items[0]['line']}")
                
                # Проверяем правильность нумерации
                expected = 1
                for item in literature_items:
                    if item['number'] != expected:
                        errors.append(f"❌ Строка {item['line']}: Неправильная нумерация! Ожидается {expected}, найдено {item['number']}")
                    expected += 1
                
                # Показываем первые 3 и последние 2 источника
                print(f"\n   ПЕРВЫЕ 3 ИСТОЧНИКА:")
                for item in literature_items[:3]:
                    print(f"      {item['number']}. {item['text']}...")
                
                print(f"\n   ПОСЛЕДНИЕ 2 ИСТОЧНИКА:")
                for item in literature_items[-2:]:
                    print(f"      {item['number']}. {item['text']}...")
    
    # Вывод ошибок
    print(f"\n{'=' * 80}")
    print(f"РЕЗУЛЬТАТЫ ПРОВЕРКИ")
    print(f"{'=' * 80}\n")
    
    if errors:
        print(f"❌ НАЙДЕНО ОШИБОК: {len(errors)}\n")
        for error in errors[:10]:  # Показываем первые 10
            print(error)
        if len(errors) > 10:
            print(f"\n... и еще {len(errors) - 10} ошибок")
    else:
        print("✅ КРИТИЧЕСКИХ ОШИБОК НЕ НАЙДЕНО!")
    
    if warnings:
        print(f"\n⚠️  ПРЕДУПРЕЖДЕНИЯ: {len(warnings)}\n")
        for warning in warnings[:5]:
            print(warning)
        if len(warnings) > 5:
            print(f"\n... и еще {len(warnings) - 5} предупреждений")
    
    # Итоговая оценка
    print(f"\n{'=' * 80}")
    if len(errors) == 0 and len(warnings) == 0:
        print("🎉 ДОКУМЕНТ ГОТОВ К ИСПОЛЬЗОВАНИЮ!")
    elif len(errors) == 0:
        print("✅ Документ в порядке, но есть незначительные замечания")
    else:
        print("⚠️  ТРЕБУЕТСЯ ДОРАБОТКА - найдены ошибки")
    print(f"{'=' * 80}\n")


if __name__ == '__main__':
    check_word_document('ДИПЛОМ_ПОЛНЫЙ.docx')
